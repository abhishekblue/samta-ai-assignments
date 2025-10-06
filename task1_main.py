# Task 1: RAG system with PDF.

import logging
from dotenv import load_dotenv
from docx import Document
from pathlib import Path
from src.config import CHUNK_SIZE, CHUNK_OVERLAP, TOP_K, MODEL, TEMPERATURE
from src.document_processor import DocumentProcessor
from src.vector_store import VectorStoreManager
from src.rag_engine import RAGEngine

load_dotenv()

# Setup logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


def main():
    # Run Task 1: PDF-based RAG system.
    logger.info("Initializing RAG system for Task 1...")
    
    doc_processor = DocumentProcessor(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP
    )
    
    # Load PDF
    pdf_path = "data/sample.pdf"  # UPDATE THIS PATH
    if not Path(pdf_path).exists():
        raise FileNotFoundError(f"PDF not found: {pdf_path}")
    
    logger.info("Loading and processing PDF...")
    documents = doc_processor.load_pdf(pdf_path)
    logger.info(f"Loaded {len(documents)} document chunks")

    # Generate Word file from extracted text
    logger.info("Generating Word file from extracted text...")
    doc = Document()
    doc.add_heading('Extracted Text from PDF', 0)
    for i, chunk in enumerate(documents):
        doc.add_heading(f'Chunk {i+1}', level=2)
        doc.add_paragraph(chunk.page_content)
    word_output_path = "data/extracted_text.docx"
    doc.save(word_output_path)
    logger.info(f"Saved extracted text to {word_output_path}")
    
    # Create vector store
    vector_manager = VectorStoreManager()
    vector_manager.create_vector_store(documents)
    retriever = vector_manager.get_retriever(k=TOP_K)
    
    # Initialize RAG engine
    rag_engine = RAGEngine(
        retriever=retriever,
        model_name=MODEL,
        temperature=TEMPERATURE
    )
    
    # Interactive query loop
    logger.info("\n" + "="*50)
    logger.info("RAG System Ready! Ask questions (type 'quit' to exit)")
    logger.info("="*50 + "\n")
    
    while True:
        question = input("\nQuestion: ").strip()
        
        if question.lower() in ['quit', 'exit', 'q']:
            break
        
        if not question:
            continue
        
        result = rag_engine.query(question)
        
        print(f"\nAnswer: {result['answer']}\n")
        print(f"Sources: {len(result['sources'])} documents used")


if __name__ == "__main__":
    main()
